from docx import Document
from docx.shared import Pt
import re

def process_word_document(input_path, output_path):
    """
    Reads a Word document, cleans up the text by removing unnecessary spaces and line breaks,
    and saves the processed document.
    """
    # Load the document
    doc = Document(input_path)
    
    # Process each paragraph in the document
    prev_paragraph = None
    for i, para in enumerate(doc.paragraphs):
        # Clean up the paragraph text
        cleaned_paragraph = re.sub(r'\n+', '\n', para.text.strip())
        
        # Remove placeholders
        cleaned_paragraph = cleaned_paragraph.replace("**** end ****", "")
        
        # Update the paragraph with the cleaned text
        para.clear()
        para.add_run(cleaned_paragraph)

        # Set space after the paragraph to 2 units
        para.paragraph_format.space_after = Pt(2)
        
        # Set space before the paragraph to 2 units
        if prev_paragraph is not None:
            prev_paragraph.paragraph_format.space_after = Pt(2)
        
        prev_paragraph = para
    
    # Save the processed document
    doc.save(output_path)
    print(f"Processed document saved as '{output_path}'")

if __name__ == "__main__":
    input_doc_path = input("Enter the path of the document to be processed: ")
    output_doc_path = input("Enter the name of the output file: ")
    process_word_document(input_doc_path, output_doc_path)